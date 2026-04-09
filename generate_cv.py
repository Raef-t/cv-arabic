import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from fpdf import FPDF
import arabic_reshaper
from bidi.algorithm import get_display

def set_rtl(paragraph):
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def set_font(run, font_name='Arial', size=11, bold=False, color=None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:hint'), 'arabic')
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def create_docx():
    doc = Document()
    
    # Modern Design: Blue Accent
    accent_color = (56, 189, 248) # Cyan-blue from web
    
    # Header
    header = doc.add_paragraph()
    set_rtl(header)
    run = header.add_run("رائف جودت")
    set_font(run, size=24, bold=True, color=accent_color)
    
    sub = doc.add_paragraph()
    set_rtl(sub)
    run = sub.add_run("مبرمج ويب & محاسب مالي")
    set_font(run, size=14, color=(100, 100, 100))
    
    # Contact Info
    contact = doc.add_paragraph()
    set_rtl(contact)
    info = "rarfjt94@gmail.com | 933219934 (963+) | حلب، سوريا"
    run = contact.add_run(info)
    set_font(run, size=10)
    
    def add_section(title):
        p = doc.add_paragraph()
        set_rtl(p)
        run = p.add_run(title)
        set_font(run, size=16, bold=True, color=accent_color)
        # Add a bottom border manually via XML is complex, so we just use styling
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)

    # Experience
    add_section("الخبرة المهنية")
    
    experiences = [
        ("مبرمج ويب - مجموعة أي إس إس (ISS Group)", "يوليو 2025 - الحالي", [
            "قيادة التطوير الكامل للواجهة الخلفية باستخدام Laravel، والتعامل مع تصميم API، وبنية قاعدة البيانات، ونشر الخادم.",
            "تحسين أداء التطبيق وضمان أمان النظام وقابليته للتوسع.",
            "تقديم ورش عمل تدريبية وتوجيه المطورين المبتدئين."
        ]),
        ("محاسب مالي - شركة اللوزي لتجارة الألبسة", "يناير 2025 - مايو 2025", [
            "تأسيس النظام المحاسبي الأساسي باستخدام Excel وبرنامج الأمين.",
            "إدارة إدخالات الفواتير والإشراف على التدفق النقدي والذمم المدينة والدائنة."
        ]),
        ("مطور ويب - المجموعة الرقمية (Digital Group)", "فبراير 2025 - مارس 2025", [
            "تطوير منصة تجارة إلكترونية متعددة البائعين ودمج بوابات الدفع.",
            "تحسين الأداء وتجربة المستخدم لعمليات التسوق."
        ]),
        ("منسق مشروع - حملة مكافحة الليشمانيا، اليونيسف (UNICEF)", "2015 - 2018", [
            "تنسيق الأنشطة الميدانية وإدارة اللوجستيات للحملة الصحية في حلب."
        ])
    ]
    
    for job, date, bullets in experiences:
        p = doc.add_paragraph()
        set_rtl(p)
        run = p.add_run(job)
        set_font(run, size=12, bold=True)
        
        d = doc.add_paragraph()
        set_rtl(d)
        run = d.add_run(date)
        set_font(run, size=9, color=(120, 120, 120))
        
        for bullet in bullets:
            bp = doc.add_paragraph(style='List Bullet')
            set_rtl(bp)
            run = bp.add_run(bullet)
            set_font(run, size=10)

    # Education
    add_section("التعليم والتدريب")
    edu_list = [
        "ماجستير في اللغويات التطبيقية | الجامعة الافتراضية السورية (2023 - الحالي)",
        "كلية الاقتصاد | جامعة حلب (2016 - 2023)",
        "دبلوم هندسة شبكات الحاسوب | جامعة حلب (2013 - 2015)",
        "الذكاء الاصطناعي وتطبيقاته في البرمجة (2024 - الحالي)"
    ]
    for edu in edu_list:
        p = doc.add_paragraph()
        set_rtl(p)
        run = p.add_run(edu)
        set_font(run, size=11)

    # Skills
    add_section("المهارات")
    skills = "PHP (Laravel), Next.js, React, C#, SQL Server, MySQL, Postman, برنامج الأمين, الإنجليزية (B2)"
    p = doc.add_paragraph()
    set_rtl(p)
    run = p.add_run(skills)
    set_font(run, size=11)

    doc.save("CV_Raef_Joudeh.docx")
    print("Word file generated successfully.")

def create_pdf():
    # Attempting to find a valid Arabic font on Windows
    font_path = r"C:\Windows\Fonts\arial.ttf" # Most common
    if not os.path.exists(font_path):
        font_path = r"C:\Windows\Fonts\tahoma.ttf"
    
    pdf = FPDF()
    pdf.add_page()
    
    try:
        pdf.add_font("Arabic", "", font_path)
        pdf.set_font("Arabic", size=16)
    except:
        print("Warning: Standard font not found, falling back to basic font (Arabic will likely break).")
        pdf.set_font("Arial", size=16)

    def write_rtl(text, size=12, bold=False, color=(0,0,0), align='R'):
        try:
            reshaped_text = arabic_reshaper.reshape(text)
            bidi_text = get_display(reshaped_text)
            pdf.set_text_color(*color)
            pdf.set_font("Arabic", size=size)
            # Using a fixed width slightly less than page width to avoid margin issues
            pdf.multi_cell(190, 10, bidi_text, align=align)
        except Exception as e:
            print(f"Error writing line: {text}. Error: {e}")

    # Header
    write_rtl("رائف جودت", size=24, color=(56, 189, 248))
    write_rtl("مبرمج ويب & محاسب مالي", size=16, color=(100, 100, 100))
    write_rtl("rarfjt94@gmail.com | 933219934 | حلب، سوريا", size=10)
    pdf.ln(5)

    # Sections - Simplified for PDF
    write_rtl("الخبرة المهنية", size=18, color=(56, 189, 248))
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(2)
    
    exp_items = [
        "مبرمج ويب - مجموعة أي إس إس (ISS Group)",
        "محاسب مالي - شركة اللوزي لتجارة الألبسة",
        "مطور ويب - المجموعة الرقمية",
        "منسق مشروع - اليونيسف UNICEF"
    ]
    for item in exp_items:
        write_rtl(f"- {item}", size=11)

    pdf.ln(5)
    write_rtl("التعليم", size=18, color=(56, 189, 248))
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(2)
    
    edu_items = [
        "ماجستير في اللغويات التطبيقية (الحالي)",
        "كلية الاقتصاد | جامعة حلب (2016-2023)",
        "دبلوم هندسة شبكات الحاسوب (2013-2015)"
    ]
    for item in edu_items:
        write_rtl(f"- {item}", size=11)

    pdf.ln(5)
    write_rtl("المهارات", size=18, color=(56, 189, 248))
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(2)
    write_rtl("PHP Laravel, Next.js, React, C#, SQL Server, برنامج الأمين", size=11)

    pdf.output("CV_Raef_Joudeh.pdf")
    print("PDF file generated successfully.")

if __name__ == "__main__":
    create_docx()
    create_pdf()
